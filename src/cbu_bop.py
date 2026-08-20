"""Parsers for CBU Balance of Payments releases.

CBU's English BOP/IIP/ED publications are year-to-date cumulative within a
calendar year, not discrete quarterly flows. The functions here (1) pull
"Secondary income, credit" out of a single release and (2) difference those
YTD figures into quarterly flows. Mixing those two steps is the easiest way
to silently triple-count Q1 by year-end.
"""

from __future__ import annotations

import re
from pathlib import Path

import pandas as pd

# v1 proxy: Addenda 1 does not break out personal transfers / compensation of
# employees. Secondary income, credit is the closest published line.
# TODO(v2): replace with IMF BOP (data.imf.org) disaggregated series —
#   personal transfers + compensation of employees — once that pull exists.
ROW_LABEL_RE = re.compile(r"secondary income,\s*credits?\b", re.I)
DEBIT_RE = re.compile(r"\bdebit", re.I)

# CBU prints millions with a space as the thousands separator and *either*
# a comma or a period as the decimal (2024+ PDFs mix both across releases).
NUMBER_RE = re.compile(
    r"[-+\u2212]?(?:\d{1,3}(?:[ \u00a0]\d{3})+|\d+)(?:[.,]\d+)?"
)
YEAR_RANGE_RE = re.compile(r"(20\d{2})\s*[-–]\s*(20\d{2})")
YEAR_RE = re.compile(r"(?<!\d)(20\d{2})(?!\d)")

# Filename patterns CBU has actually used, oldest to newest. Order matters:
# YYYYQn must win over a bare year, or 2026Q1 would be read as an annual file.
_PERIOD_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"(?P<year>20\d{2})Q(?P<q>[1-4])", re.I), "yq"),
    (re.compile(r"(?P<year>20\d{2})H1", re.I), "h1_year"),
    (re.compile(r"9M[_-]?(?P<year>20\d{2})", re.I), "q3"),
    (re.compile(r"1H[_-]?(?P<year>20\d{2})", re.I), "q2"),
    (re.compile(r"(?P<q>[1-4])Q[_-]?(?P<year>20\d{2})", re.I), "qy"),
    (re.compile(r"(?<![QqHhy\d])(?P<year>20\d{2})(?!Q|H|\d)", re.I), "annual"),
]


class CbuBopParseError(ValueError):
    """A CBU release could not be mapped to a (year, quarter, YTD value)."""


def parse_cbu_number(token: str) -> float:
    """Parse a CBU-formatted million-USD figure ('1 690.1' or '6 957,2')."""
    s = token.strip().replace("\u00a0", " ").replace(" ", "")
    s = s.replace("\u2212", "-")
    if "," in s and "." in s:
        # Last separator is the decimal; the other is thousands.
        if s.rfind(",") > s.rfind("."):
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", "")
    elif "," in s:
        s = s.replace(",", ".")
    return float(s)


def parse_release_period(filename: str) -> tuple[int, int]:
    """Infer (year, quarter) from a CBU BOP filename.

    Q2 = first-half / H1 release (cumulative through June).
    Q3 = 9-month release. Q4 = full-year release. A bare year with no
    quarter marker is treated as Q4 (the annual publication).
    """
    stem = Path(filename).stem
    for pattern, kind in _PERIOD_PATTERNS:
        m = pattern.search(stem)
        if not m:
            continue
        year = int(m.group("year"))
        if kind == "yq":
            return year, int(m.group("q"))
        if kind == "qy":
            return year, int(m.group("q"))
        if kind in {"h1_year", "q2"}:
            return year, 2
        if kind == "q3":
            return year, 3
        if kind == "annual":
            print(
                f"[cbu-bop] {filename}: no quarter in the name; treating as "
                f"{year}Q4 (full-year YTD). Rename to eng_BOP_IIP_ED_{year}Q4.* "
                "if that is wrong."
            )
            return year, 4
    raise CbuBopParseError(
        f"Cannot infer year/quarter from {filename!r}. Expected names like "
        "eng_BOP_IIP_ED_2025Q3.pdf, eng_BOP_IIP_9M_2023.docx, or "
        "eng_BOP_IIP_ED_2024H1.pdf. Rename BALANCE-OF-PAYMENTS_*.docx to "
        "eng_BOP_IIP_ED_2019Q4.docx (or whatever year it is) before parsing."
    )


def _unique_keep_order(xs: list[int]) -> list[int]:
    seen: set[int] = set()
    out: list[int] = []
    for x in xs:
        if x not in seen:
            seen.add(x)
            out.append(x)
    return out


def _years_from_text(text: str) -> list[int]:
    years: list[int] = []
    consumed: list[tuple[int, int]] = []
    for m in YEAR_RANGE_RE.finditer(text):
        y0, y1 = int(m.group(1)), int(m.group(2))
        if y1 >= y0 and (y1 - y0) <= 5:
            years.extend(range(y0, y1 + 1))
            consumed.append(m.span())
    for m in YEAR_RE.finditer(text):
        if any(a <= m.start() < b for a, b in consumed):
            continue
        years.append(int(m.group(1)))
    return _unique_keep_order(years)


def parse_secondary_income_from_text(text: str, reporting_year: int) -> float:
    """Pull Secondary income, credit for `reporting_year` out of Addenda 1 text.

    Recent CBU PDFs split the analytic table into fragments that pdfplumber
    cannot reassemble, so we parse the extracted text line rather than the
    table object. The rightmost column is the reporting year when headers
    cannot be aligned — that is logged, never silently guessed past.
    """
    text = text.replace("\u2212", "-")
    lines = text.splitlines()
    value_line_idx: int | None = None
    numbers: list[float] | None = None
    for i, line in enumerate(lines):
        if not ROW_LABEL_RE.search(line) or DEBIT_RE.search(line):
            continue
        found = [parse_cbu_number(m.group(0)) for m in NUMBER_RE.finditer(line)]
        # The label itself contains no years; every hit should be a cell value.
        if found:
            value_line_idx = i
            numbers = found
            break
    if numbers is None or value_line_idx is None:
        raise CbuBopParseError(
            "No 'Secondary income, credit' line found. This parser looks at "
            "Addenda 1 (analytic presentation), not the narrative charts."
        )

    header_window = "\n".join(lines[max(0, value_line_idx - 12) : value_line_idx])
    years = _years_from_text(header_window)
    if reporting_year in years and len(years) == len(numbers):
        return numbers[years.index(reporting_year)]
    if reporting_year in years and len(years) != len(numbers):
        print(
            f"[cbu-bop] header years {years} do not line up with "
            f"{len(numbers)} value(s); falling back to the rightmost column "
            f"for {reporting_year}."
        )
    else:
        print(
            f"[cbu-bop] could not match {reporting_year} to header years "
            f"{years or '(none)'}; using the rightmost Addenda 1 column. "
            "Check this against the PDF if the figure looks off."
        )
    return numbers[-1]


def extract_secondary_income_credit(path: Path, reporting_year: int) -> float:
    """Read one CBU BOP release and return Secondary income, credit (mn USD).

    2019–2023 releases have .docx twins that python-docx can table-parse
    cleanly; 2024+ is PDF-only and goes through pdfplumber text extraction.
    """
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _extract_from_docx(path, reporting_year)
    if suffix == ".pdf":
        return _extract_from_pdf(path, reporting_year)
    raise CbuBopParseError(f"Unsupported CBU BOP file type: {path.name}")


def _extract_from_docx(path: Path, reporting_year: int) -> float:
    from docx import Document  # imported here so PDF-only runs don't need it

    doc = Document(str(path))
    for table in doc.tables:
        if not table.rows:
            continue
        header = [c.text.replace("\n", " ") for c in table.rows[0].cells]
        years = _years_from_text(" ".join(header))
        for row in table.rows[1:]:
            cells = [c.text.replace("\n", " ").strip() for c in row.cells]
            if not cells or not ROW_LABEL_RE.search(cells[0]) or DEBIT_RE.search(cells[0]):
                continue
            values = [parse_cbu_number(c) for c in cells[1:] if NUMBER_RE.search(c)]
            if not values:
                continue
            if reporting_year in years and len(years) == len(values):
                value = values[years.index(reporting_year)]
            else:
                value = values[-1]
                print(
                    f"[cbu-bop] {path.name}: header years {years} vs "
                    f"{len(values)} value(s); using rightmost column "
                    f"({value}) for {reporting_year}."
                )
            print(f"[cbu-bop] {path.name}: Secondary income, credit = {value} mn USD (YTD)")
            return value
    # Some older files bury the table in a way python-docx misses.
    text = "\n".join(p.text for p in doc.paragraphs)
    value = parse_secondary_income_from_text(text, reporting_year)
    print(f"[cbu-bop] {path.name}: Secondary income, credit = {value} mn USD (YTD, paragraph fallback)")
    return value


def _extract_from_pdf(path: Path, reporting_year: int) -> float:
    import pdfplumber

    with pdfplumber.open(path) as pdf:
        text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    value = parse_secondary_income_from_text(text, reporting_year)
    print(f"[cbu-bop] {path.name}: Secondary income, credit = {value} mn USD (YTD)")
    return value


def ytd_to_quarterly(ytd: pd.Series) -> pd.Series:
    """Convert within-year YTD cumulatives into discrete quarterly flows.

    CRITICAL — do not skip this step and do not treat a missing predecessor
    as zero. CBU Q1 = Q1 flow, Q2 publication = H1 = Q1+Q2, Q3 = 9 months,
    Q4 = full year. Reading each release as-is would count Q1 four times by
    December and fabricate a unit-root / cointegration result out of an
    accounting identity.

        Q1 flow = YTD_Q1
        Q2 flow = YTD_H1  - YTD_Q1
        Q3 flow = YTD_9M  - YTD_H1
        Q4 flow = YTD_FY  - YTD_9M

    YTD resets on January 1, so 2025Q1 is NOT differenced against 2024Q4.
    If Qn exists but Q(n-1) does not, that quarter is NaN and a warning is
    printed — using the YTD figure as if it were the quarterly flow would
    be the same class of error as skipping the difference entirely.
    """
    if ytd.empty:
        raise ValueError("ytd_to_quarterly() got an empty series.")

    s = ytd.copy()
    if isinstance(s.index, pd.DatetimeIndex):
        s.index = s.index.to_period("Q")
    elif not isinstance(s.index, pd.PeriodIndex):
        s.index = pd.PeriodIndex(s.index, freq="Q")
    s = s.sort_index()
    s = pd.to_numeric(s, errors="coerce")

    records: list[tuple[pd.Period, float]] = []
    for year, grp in s.groupby(s.index.year):
        ytd_q = {int(period.quarter): float(val) for period, val in grp.items()}
        for q in range(1, 5):
            if q not in ytd_q:
                continue
            period = pd.Period(year=int(year), quarter=q, freq="Q")
            if q == 1:
                flow = ytd_q[1]
                print(f"[cbu-bop] {period} flow = Q1 YTD {flow:.1f} (no difference)")
            else:
                prev = ytd_q.get(q - 1)
                if prev is None or pd.isna(prev):
                    print(
                        f"[cbu-bop] WARNING: {period} YTD={ytd_q[q]:.1f} exists but "
                        f"{year}Q{q - 1} does not. Cannot difference without inventing "
                        f"the missing predecessor. Leaving {period} as NaN — not using "
                        f"the YTD figure as a quarterly flow."
                    )
                    flow = float("nan")
                else:
                    flow = ytd_q[q] - prev
                    print(
                        f"[cbu-bop] {period} flow = {year}Q{q} YTD {ytd_q[q]:.1f} "
                        f"- {year}Q{q - 1} YTD {prev:.1f} = {flow:.1f}"
                    )
                    if flow < 0:
                        print(
                            f"[cbu-bop] WARNING: {period} flow is negative ({flow:.1f}). "
                            "Could be a CBU vintage revision (later YTD < earlier YTD) "
                            "or a parse error. Not dropping the observation."
                        )
            records.append((period, flow))

    out = pd.Series(
        {period: value for period, value in records},
        dtype="float64",
        name="remittances_mn_usd",
    )
    out.index = pd.PeriodIndex(out.index, freq="Q")
    return out


def select_preferred_releases(paths: list[Path]) -> dict[tuple[int, int], Path]:
    """One file per (year, quarter). Prefer .docx over .pdf when both exist."""
    chosen: dict[tuple[int, int], Path] = {}
    for path in sorted(paths):
        year, quarter = parse_release_period(path.name)
        key = (year, quarter)
        existing = chosen.get(key)
        if existing is None:
            chosen[key] = path
            continue
        if path.suffix.lower() == ".docx" and existing.suffix.lower() != ".docx":
            print(
                f"[cbu-bop] {year}Q{quarter}: preferring {path.name} over "
                f"{existing.name} (.docx is easier to parse than the PDF twin)."
            )
            chosen[key] = path
        elif existing.suffix.lower() == ".docx" and path.suffix.lower() != ".docx":
            print(
                f"[cbu-bop] {year}Q{quarter}: keeping {existing.name}, ignoring {path.name}."
            )
        else:
            print(
                f"[cbu-bop] WARNING: {year}Q{quarter} has both {existing.name} and "
                f"{path.name}; keeping {existing.name}. Delete one if that is wrong."
            )
    return chosen
